from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy import select

from app.core.database import AsyncSessionLocal
from app.models.session import Session
from app.models.message import Message, MessageType, MessageRole
from app.models.media import Media


class DocumentGenerator:
    async def generate_requirements_document(
        self,
        session: Session,
        format: str = "docx",
    ) -> bytes:
        """Generate a requirements document from session data."""
        if format == "docx":
            return await self._generate_docx(session)
        else:
            return await self._generate_markdown(session)

    async def _generate_docx(self, session: Session) -> bytes:
        doc = Document()

        # Title
        title = doc.add_heading(session.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
        if session.description:
            doc.add_paragraph(f"Description: {session.description}")
        doc.add_paragraph()

        # Table of Contents placeholder
        doc.add_heading("Table of Contents", level=1)
        doc.add_paragraph("(Table of contents will be generated in Word)")
        doc.add_page_break()

        # Extract requirements from messages
        async with AsyncSessionLocal() as db:
            result = await db.execute(
                select(Message)
                .where(Message.session_id == session.id)
                .order_by(Message.created_at.asc())
            )
            messages = result.scalars().all()

        # Requirements section
        doc.add_heading("Requirements", level=1)

        requirement_count = 0
        for msg in messages:
            if msg.message_type == MessageType.REQUIREMENT:
                requirement_count += 1
                doc.add_heading(f"REQ-{requirement_count:03d}", level=2)
                doc.add_paragraph(msg.content)
                if msg.extra_data:
                    if msg.extra_data.get("category"):
                        doc.add_paragraph(f"Category: {msg.extra_data['category']}")
                    if msg.extra_data.get("priority"):
                        doc.add_paragraph(f"Priority: {msg.extra_data['priority']}")
                doc.add_paragraph()

        if requirement_count == 0:
            doc.add_paragraph("No formal requirements have been captured yet.")

        # Conversation Summary section
        doc.add_page_break()
        doc.add_heading("Conversation Summary", level=1)

        for msg in messages:
            if msg.message_type == MessageType.TEXT and msg.role == MessageRole.ASSISTANT:
                # Add key assistant messages as summary points
                if len(msg.content) > 50:  # Skip very short messages
                    p = doc.add_paragraph()
                    p.add_run(f"• {msg.content[:500]}...")
                    if len(msg.content) > 500:
                        p.add_run(" [truncated]")

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    async def _generate_markdown(self, session: Session) -> bytes:
        lines = []

        # Title
        lines.append(f"# {session.title}")
        lines.append("")
        lines.append(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
        if session.description:
            lines.append(f"Description: {session.description}")
        lines.append("")

        # Extract requirements from messages
        async with AsyncSessionLocal() as db:
            result = await db.execute(
                select(Message)
                .where(Message.session_id == session.id)
                .order_by(Message.created_at.asc())
            )
            messages = result.scalars().all()

        # Requirements section
        lines.append("## Requirements")
        lines.append("")

        requirement_count = 0
        for msg in messages:
            if msg.message_type == MessageType.REQUIREMENT:
                requirement_count += 1
                lines.append(f"### REQ-{requirement_count:03d}")
                lines.append("")
                lines.append(msg.content)
                if msg.extra_data:
                    if msg.extra_data.get("category"):
                        lines.append(f"- **Category:** {msg.extra_data['category']}")
                    if msg.extra_data.get("priority"):
                        lines.append(f"- **Priority:** {msg.extra_data['priority']}")
                lines.append("")

        if requirement_count == 0:
            lines.append("No formal requirements have been captured yet.")
            lines.append("")

        return "\n".join(lines).encode("utf-8")

    async def generate_session_summary(self, session: Session) -> str:
        """Generate a text summary of the session."""
        async with AsyncSessionLocal() as db:
            result = await db.execute(
                select(Message)
                .where(Message.session_id == session.id)
                .order_by(Message.created_at.asc())
            )
            messages = result.scalars().all()

        total_messages = len(messages)
        user_messages = sum(1 for m in messages if m.role == MessageRole.USER)
        assistant_messages = sum(1 for m in messages if m.role == MessageRole.ASSISTANT)
        requirements = sum(1 for m in messages if m.message_type == MessageType.REQUIREMENT)
        questionnaires = sum(1 for m in messages if m.message_type == MessageType.QUESTIONNAIRE)

        summary = f"""Session Summary: {session.title}

Status: {session.status.value}
Created: {session.created_at.strftime('%Y-%m-%d %H:%M UTC')}
Last Updated: {session.updated_at.strftime('%Y-%m-%d %H:%M UTC')}

Statistics:
- Total Messages: {total_messages}
- User Messages: {user_messages}
- Assistant Messages: {assistant_messages}
- Requirements Captured: {requirements}
- Questionnaires Generated: {questionnaires}
- Token Usage: {session.token_usage}
"""
        return summary

    async def export_session(
        self,
        session: Session,
        include_messages: bool = True,
        include_media: bool = False,
    ) -> dict:
        """Export session data as JSON."""
        export_data = {
            "id": str(session.id),
            "title": session.title,
            "description": session.description,
            "status": session.status.value,
            "created_at": session.created_at.isoformat(),
            "updated_at": session.updated_at.isoformat(),
            "token_usage": session.token_usage,
        }

        async with AsyncSessionLocal() as db:
            if include_messages:
                result = await db.execute(
                    select(Message)
                    .where(Message.session_id == session.id)
                    .order_by(Message.created_at.asc())
                )
                messages = result.scalars().all()
                export_data["messages"] = [
                    {
                        "id": str(m.id),
                        "role": m.role.value,
                        "message_type": m.message_type.value,
                        "content": m.content,
                        "metadata": m.metadata,
                        "created_at": m.created_at.isoformat(),
                    }
                    for m in messages
                ]

            if include_media:
                result = await db.execute(
                    select(Media)
                    .where(Media.session_id == session.id)
                    .order_by(Media.created_at.asc())
                )
                media_files = result.scalars().all()
                export_data["media"] = [
                    {
                        "id": str(m.id),
                        "filename": m.original_filename,
                        "content_type": m.content_type,
                        "media_type": m.media_type.value,
                        "size_bytes": m.size_bytes,
                        "url": m.storage_url,
                        "created_at": m.created_at.isoformat(),
                    }
                    for m in media_files
                ]

        return export_data
